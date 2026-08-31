from __future__ import annotations

import csv
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"
NS = {"w": W_NS, "w14": W14_NS}


def write_docx(
    project: Path,
    paragraphs: list[str],
    table_rows: list[list[str]] | None = None,
    *,
    unsupported: bool = False,
) -> Path:
    path = project / "paper" / "final.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
    document.save(path)
    if unsupported:
        _inject_unsupported_drawing(path)
    return path


def _inject_unsupported_drawing(path: Path) -> None:
    """Add a real OOXML drawing marker while retaining a valid DOCX fixture."""

    with zipfile.ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}
    root = etree.fromstring(members["word/document.xml"])
    body = root.find("w:body", namespaces=NS)
    assert body is not None
    paragraph = etree.Element(f"{{{W_NS}}}p", nsmap={"w": W_NS, "w14": W14_NS})
    run = etree.SubElement(paragraph, f"{{{W_NS}}}r")
    etree.SubElement(run, f"{{{W_NS}}}drawing")
    body.insert(max(0, len(body) - 1), paragraph)
    members["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    temporary = path.with_suffix(".patched.docx")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as destination:
        for name, content in members.items():
            destination.writestr(name, content)
    temporary.replace(path)


def inject_revision_insert(path: Path, text: str) -> None:
    """Insert visible text under a real w:ins revision container."""

    with zipfile.ZipFile(path, "r") as source:
        members = {name: source.read(name) for name in source.namelist()}
    root = etree.fromstring(members["word/document.xml"])
    body = root.find("w:body", namespaces=NS)
    assert body is not None
    paragraph = etree.Element(f"{{{W_NS}}}p", nsmap={"w": W_NS, "w14": W14_NS})
    inserted = etree.SubElement(paragraph, f"{{{W_NS}}}ins")
    run = etree.SubElement(inserted, f"{{{W_NS}}}r")
    text_node = etree.SubElement(run, f"{{{W_NS}}}t")
    text_node.text = text
    body.insert(max(0, len(body) - 1), paragraph)
    members["word/document.xml"] = etree.tostring(
        root, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    temporary = path.with_suffix(".revision.docx")
    with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as destination:
        for name, content in members.items():
            destination.writestr(name, content)
    temporary.replace(path)


def write_config(project: Path) -> Path:
    path = project / "mea.toml"
    path.write_text(
        "schema_version = 1\n"
        'manuscript = "paper/final.docx"\n'
        'registry = "state/frozen_numbers.csv"\n'
        'mapping = "state/claim_map.csv"\n'
        'output_dir = "build"\n'
        "ignore_years = true\n"
        "require_frozen = true\n",
        encoding="utf-8",
        newline="",
    )
    return path


def write_csv(path: Path, headers: list[str], rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=headers, lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)


REGISTRY_HEADERS = [
    "claim_id",
    "question",
    "claim_text",
    "metric",
    "value",
    "unit",
    "source_file",
    "source_column",
    "filter",
    "run_id",
    "status",
    "display_value",
    "evidence_type",
    "round_digits",
    "tolerance_abs",
    "tolerance_rel",
    "notes",
]


def write_registry(project: Path, rows: list[dict[str, str]]) -> Path:
    return _write_and_return(project / "state" / "frozen_numbers.csv", REGISTRY_HEADERS, rows)


def write_mapping(project: Path, rows: list[dict[str, str]]) -> Path:
    return _write_and_return(
        project / "state" / "claim_map.csv",
        ["occurrence_id", "claim_id", "decision", "context", "confirmed_at"],
        rows,
    )


def write_evidence(project: Path, rows: list[dict[str, str]]) -> Path:
    return _write_and_return(
        project / "results" / "metrics.csv",
        ["model", "split", "metric", "value", "run_id"],
        rows,
    )


def _write_and_return(path: Path, headers: list[str], rows: list[dict[str, str]]) -> Path:
    write_csv(path, headers, rows)
    return path


def make_project(
    tmp_path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None
) -> Path:
    project = tmp_path / "中文项目"
    project.mkdir()
    write_docx(project, paragraphs, table_rows)
    write_config(project)
    write_mapping(project, [])
    return project
